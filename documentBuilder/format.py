import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from ebooklib import epub
import subprocess

def export_document(content, export_format='pdf', output_file='document'):
    """
    Export the given content to a specified format.

    Args:
    - content: The text content (str) to be exported (can include title and abstract).
    - export_format: The format to export to ('pdf', 'docx', 'epub', 'txt').
    - output_file: The output filename (without extension).

    Supported formats: 'pdf', 'docx', 'epub', 'txt'.
    """
    
    # if export_format == 'pdf':
    #     # Export to PDF using ReportLab
    #     export_to_pdf(content, output_file)
    if export_format == 'docx':
        # Export to DOCX using python-docx
        export_to_docx(content, output_file)
    # elif export_format == 'epub':
    #     # Export to EPUB using EbookLib
    #     export_to_epub(content, output_file)
    # elif export_format == 'txt':
    #     # Export to TXT using standard file writing
    #     export_to_txt(content, output_file)
    # elif export_format == 'latex':
    #     export_using_pandoc(content, output_file, 'tex')
    # elif export_format == 'odt':
    #     export_using_pandoc(content, output_file, 'odt')
    # elif export_format == 'html':
    #     export_to_html(content, output_file)
    # elif export_format == 'mobi':
    #     export_using_calibre(content, output_file, 'mobi')
    else:
        print(f"Unsupported format: {export_format}")
        return

# def export_to_pdf(content, output_file):
#     """
#     Export content to a PDF file using ReportLab.
#     """
#     c = canvas.Canvas(f"{output_file}.pdf", pagesize=letter)
#     text_object = c.beginText(100, 750)
#     text_object.setFont("Helvetica", 12)
#     text_object.textLines(content)
#     c.drawText(text_object)
#     c.save()
#     print(f"PDF saved as {output_file}.pdf")

def export_to_docx(content, output_file):
    """
    Export content to a DOCX file using python-docx.
    """
    doc = Document()
    doc.add_heading('Document', 0)
    doc.add_paragraph(content)
    doc.save(f"{output_file}.docx")
    print(f"DOCX saved as {output_file}.docx")

# def export_to_epub(content, output_file):
#     """
#     Export content to an EPUB file using EbookLib.
#     """
#     book = epub.EpubBook()
#     book.set_title('Generated Document')
#     book.set_language('en')

#     # Create a chapter for the content
#     chapter = epub.EpubHtml(title='Abstract', file_name='chapter.xhtml', lang='en')
#     chapter.content = f"<html><body><h1>Abstract</h1><p>{content}</p></body></html>"
#     book.add_item(chapter)

#     # Define the spine
#     book.spine = ['nav', chapter]

#     # Save the book as EPUB
#     epub.write_epub(f"{output_file}.epub", book)
#     print(f"EPUB saved as {output_file}.epub")

# def export_to_txt(content, output_file):
#     """
#     Export content to a TXT file.
#     """
#     with open(f"{output_file}.txt", 'w') as file:
#         file.write(content)
#     print(f"TXT saved as {output_file}.txt")

# # New export function examples:
# def export_using_calibre(content, output_file, format_type):
#     """
#     Export content to a specified format using Calibre (e.g., MOBI, AZW3).
#     """
#     # This assumes you have Calibre installed and its `ebook-convert` command-line tool available
#     with open(f"{output_file}.txt", 'w') as file:
#         file.write(content)
#     subprocess.run(['ebook-convert', f"{output_file}.txt", f"{output_file}.{format_type}"])
#     print(f"File saved as {output_file}.{format_type}")

# def export_to_html(content, output_file):
#     """
#     Export content to an HTML file.
#     """
#     with open(f"{output_file}.html", 'w') as file:
#         file.write(f"<html><body><h1>Document</h1><p>{content}</p></body></html>")
#     print(f"HTML saved as {output_file}.html")

# def convert_using_pandoc(input_file, output_format):
#     """
#     Use Pandoc to convert between various formats.
#     """
#     try:
#         subprocess.run(['pandoc', input_file, '-o', f'{os.path.splitext(input_file)[0]}.{output_format}'], check=True)
#         print(f"Converted {input_file} to {output_format}")
#     except subprocess.CalledProcessError as e:
#         print(f"Error converting file using Pandoc: {e}")

# def export_using_pandoc(content, output_file, format_type):
#     """
#     Use Pandoc to convert content to the specified format.
#     Supported formats: odt, rtf, latex (tex).
#     """
#     # Save the content temporarily as a markdown file
#     with open(f"{output_file}.md", 'w') as file:
#         file.write(content)

#     # Use Pandoc to convert the markdown file to the desired format
#     subprocess.run(['pandoc', f"{output_file}.md", '-o', f"{output_file}.{format_type}"])
#     print(f"File saved as {output_file}.{format_type}")

# Example usage
content = """
Title: Example of Abstract Generation

Abstract:
This document demonstrates how to generate a file with an abstract that can be exported to multiple formats such as EPUB, DOCX, PDF, and TXT. The process involves using various tools that can export to these formats.
"""

# Export to DOCX
export_document(content, 'docx', 'document')